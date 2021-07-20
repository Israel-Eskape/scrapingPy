from bs4 import BeautifulSoup
import requests
import pandas as pd 
from docx import Document

print("Name of band : ")
band = input()
print("Name of song : ")
song = input()

band = band.replace(' ','_')
song = song.replace(' ','_')

url = 'https://acordes.lacuerda.net/'+band+'/'+song+'.shtml'
page = requests.get(url)
soup = BeautifulSoup(page.content, 'html.parser')

le = soup.find_all('div', id= 't_body')
name = soup.find_all('div', id = 'tH1')

names = list()
letra = list()
for j in name:
    names.append(j.text)

for i in le:
   letra.append(i.text)

if(len(letra) != 0):
    names = "".join(names)
    letra  = "".join(letra)
    print(names)
    print(letra)

    print("Save as document ? Y/N")
    resp = input()
    if(resp == 'Y' or resp == 'y' or resp == 'S' or resp == 's'):
        print(' saving ... ')
        document = Document()
        document.add_heading(names,0)
        document.add_paragraph(letra)

        document. save(song+'.docx')
        print(song+'.docx \n saved ... ')
else:
    print('No se encontró la canción intente de nuevo')
